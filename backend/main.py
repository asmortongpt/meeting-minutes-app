"""
Meeting Minutes API Backend
"""
from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime, date
import sqlite3
import json
from pathlib import Path
import os
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import AI analyzer
from ai_analyzer import (
    analyze_screenshot,
    analyze_meeting_content,
    identify_speakers,
    crop_and_enhance_screenshot
)

app = FastAPI(title="Meeting Minutes API")

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Database setup
DB_PATH = Path(__file__).parent / "meetings.db"
TEMPLATE_PATH = Path("/Users/andrewmorton/Downloads/Meeting Agenda Template v1- short form.docx")
EXPORT_DIR = Path(__file__).parent / "exports"
EXPORT_DIR.mkdir(exist_ok=True)
UPLOAD_DIR = Path(__file__).parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

def init_db():
    """Initialize database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS meetings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_name TEXT NOT NULL,
            meeting_date TEXT NOT NULL,
            meeting_purpose TEXT,
            agenda_items TEXT,
            attendees TEXT,
            action_items TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
    """)

    conn.commit()
    conn.close()

init_db()

# Pydantic models
class AgendaItem(BaseModel):
    item: str
    notes: str

class Attendee(BaseModel):
    name: str
    attended: bool

class ActionItem(BaseModel):
    description: str
    owner: str
    due_date: Optional[str] = None
    status: str = "Pending"

class MeetingMinutes(BaseModel):
    project_name: str
    meeting_date: str
    meeting_purpose: str
    agenda_items: List[AgendaItem]
    attendees: List[Attendee]
    action_items: List[ActionItem]

class MeetingMinutesResponse(MeetingMinutes):
    id: int
    created_at: str
    updated_at: str

@app.get("/")
def read_root():
    return {"message": "Meeting Minutes API", "version": "1.0.0"}

@app.get("/api/meetings", response_model=List[MeetingMinutesResponse])
def get_meetings():
    """Get all meetings"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("SELECT * FROM meetings ORDER BY meeting_date DESC")
    rows = cursor.fetchall()
    conn.close()

    meetings = []
    for row in rows:
        meeting = dict(row)
        meeting['agenda_items'] = json.loads(meeting['agenda_items'])
        meeting['attendees'] = json.loads(meeting['attendees'])
        meeting['action_items'] = json.loads(meeting['action_items'])
        meetings.append(meeting)

    return meetings

@app.get("/api/meetings/{meeting_id}", response_model=MeetingMinutesResponse)
def get_meeting(meeting_id: int):
    """Get a specific meeting"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("SELECT * FROM meetings WHERE id = ?", (meeting_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Meeting not found")

    meeting = dict(row)
    meeting['agenda_items'] = json.loads(meeting['agenda_items'])
    meeting['attendees'] = json.loads(meeting['attendees'])
    meeting['action_items'] = json.loads(meeting['action_items'])

    return meeting

@app.post("/api/meetings", response_model=MeetingMinutesResponse)
def create_meeting(meeting: MeetingMinutes):
    """Create a new meeting"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    now = datetime.now().isoformat()

    cursor.execute("""
        INSERT INTO meetings
        (project_name, meeting_date, meeting_purpose, agenda_items, attendees, action_items, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        meeting.project_name,
        meeting.meeting_date,
        meeting.meeting_purpose,
        json.dumps([item.dict() for item in meeting.agenda_items]),
        json.dumps([att.dict() for att in meeting.attendees]),
        json.dumps([action.dict() for action in meeting.action_items]),
        now,
        now
    ))

    meeting_id = cursor.lastrowid
    conn.commit()
    conn.close()

    return get_meeting(meeting_id)

@app.put("/api/meetings/{meeting_id}", response_model=MeetingMinutesResponse)
def update_meeting(meeting_id: int, meeting: MeetingMinutes):
    """Update a meeting"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    now = datetime.now().isoformat()

    cursor.execute("""
        UPDATE meetings
        SET project_name = ?, meeting_date = ?, meeting_purpose = ?,
            agenda_items = ?, attendees = ?, action_items = ?, updated_at = ?
        WHERE id = ?
    """, (
        meeting.project_name,
        meeting.meeting_date,
        meeting.meeting_purpose,
        json.dumps([item.dict() for item in meeting.agenda_items]),
        json.dumps([att.dict() for att in meeting.attendees]),
        json.dumps([action.dict() for action in meeting.action_items]),
        now,
        meeting_id
    ))

    if cursor.rowcount == 0:
        conn.close()
        raise HTTPException(status_code=404, detail="Meeting not found")

    conn.commit()
    conn.close()

    return get_meeting(meeting_id)

@app.delete("/api/meetings/{meeting_id}")
def delete_meeting(meeting_id: int):
    """Delete a meeting"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    cursor.execute("DELETE FROM meetings WHERE id = ?", (meeting_id,))

    if cursor.rowcount == 0:
        conn.close()
        raise HTTPException(status_code=404, detail="Meeting not found")

    conn.commit()
    conn.close()

    return {"message": "Meeting deleted successfully"}

@app.get("/api/meetings/{meeting_id}/export")
def export_meeting(meeting_id: int):
    """Export meeting to DOCX using template"""
    meeting_data = get_meeting(meeting_id)

    if not TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail="Template file not found")

    # Load template
    doc = Document(str(TEMPLATE_PATH))

    # Fill in project name and date (Table 1)
    table1 = doc.tables[0]
    table1.rows[1].cells[0].text = meeting_data['project_name']
    table1.rows[1].cells[1].text = meeting_data['meeting_date']

    # Fill in meeting purpose (Table 2)
    table2 = doc.tables[1]
    table2.rows[1].cells[0].text = meeting_data['meeting_purpose']

    # Fill in agenda items (Table 3)
    table3 = doc.tables[2]
    # Remove empty rows and keep header
    while len(table3.rows) > 2:
        table3._element.remove(table3.rows[-1]._element)

    # Add agenda items
    for agenda_item in meeting_data['agenda_items']:
        row = table3.add_row()
        row.cells[0].text = agenda_item['item']
        row.cells[1].text = agenda_item['notes']

    # Fill in attendees (Table 4)
    table4 = doc.tables[3]
    # Remove empty rows and keep header
    while len(table4.rows) > 2:
        table4._element.remove(table4.rows[-1]._element)

    # Add attendees (2 per row)
    attendees = meeting_data['attendees']
    for i in range(0, len(attendees), 2):
        row = table4.add_row()
        # First attendee
        row.cells[0].text = attendees[i]['name']
        row.cells[1].text = "✓" if attendees[i]['attended'] else ""
        # Second attendee if exists
        if i + 1 < len(attendees):
            row.cells[2].text = attendees[i + 1]['name']
            row.cells[3].text = "✓" if attendees[i + 1]['attended'] else ""

    # Fill in action items (Table 5)
    table5 = doc.tables[4]
    # Remove empty rows and keep header
    while len(table5.rows) > 2:
        table5._element.remove(table5.rows[-1]._element)

    # Add action items
    for action_item in meeting_data['action_items']:
        row = table5.add_row()
        row.cells[0].text = action_item['description']
        row.cells[1].text = action_item['owner']
        row.cells[2].text = action_item['due_date'] or ""
        row.cells[3].text = action_item['status']

    # Save document
    output_filename = f"meeting_minutes_{meeting_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = EXPORT_DIR / output_filename
    doc.save(str(output_path))

    return FileResponse(
        path=output_path,
        filename=output_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ============================================================================
# AI-POWERED ENDPOINTS (OPTIONAL - ONLY USE WHEN NEEDED TO SAVE TOKENS)
# ============================================================================

@app.post("/api/screenshots/upload")
async def upload_screenshot(file: UploadFile = File(...)):
    """Upload a screenshot without AI analysis - just stores it"""
    try:
        # Save file
        file_path = UPLOAD_DIR / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file.filename}"
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        return {
            "success": True,
            "file_path": str(file_path),
            "filename": file.filename
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/screenshots/analyze")
async def analyze_screenshot_endpoint(
    file_path: str,
    context: Optional[str] = ""
):
    """
    AI Analysis (COSTS TOKENS) - Only use when you need AI to extract info from screenshot
    Analyzes screenshot and extracts meeting info using Claude Vision
    """
    if not Path(file_path).exists():
        raise HTTPException(status_code=404, detail="Screenshot not found")

    result = analyze_screenshot(file_path, context)
    return result

@app.post("/api/screenshots/identify-speakers")
async def identify_speakers_endpoint(file_path: str):
    """
    AI Analysis (COSTS TOKENS) - Only use when you need to identify speakers
    Identifies speakers from a meeting screenshot
    """
    if not Path(file_path).exists():
        raise HTTPException(status_code=404, detail="Screenshot not found")

    speakers = identify_speakers(file_path)
    return {"speakers": speakers}

@app.post("/api/screenshots/crop")
async def crop_screenshot_endpoint(
    file_path: str,
    x: Optional[int] = None,
    y: Optional[int] = None,
    width: Optional[int] = None,
    height: Optional[int] = None
):
    """Crop and enhance screenshot (no AI, no token cost)"""
    if not Path(file_path).exists():
        raise HTTPException(status_code=404, detail="Screenshot not found")

    crop_coords = None
    if all([x is not None, y is not None, width, height]):
        crop_coords = {"x": x, "y": y, "width": width, "height": height}

    processed_path = crop_and_enhance_screenshot(file_path, crop_coords)
    return {"processed_path": processed_path}

@app.post("/api/meetings/ai-generate")
async def generate_meeting_minutes(
    meeting_notes: str,
    screenshot_paths: Optional[List[str]] = None,
    additional_context: Optional[str] = ""
):
    """
    AI Generation (COSTS TOKENS) - Only use when you want AI to auto-generate meeting minutes

    This endpoint uses Claude to automatically generate structured meeting minutes
    from your notes and optionally from screenshot analysis.

    USE SPARINGLY TO SAVE TOKENS!

    Returns structured data you can use to create a meeting record.
    """
    # Optionally analyze screenshots first if provided
    screenshots_analysis = []
    if screenshot_paths:
        for path in screenshot_paths:
            if Path(path).exists():
                analysis = analyze_screenshot(path)
                screenshots_analysis.append(analysis)

    # Generate meeting minutes
    result = analyze_meeting_content(
        meeting_notes=meeting_notes,
        screenshots_analysis=screenshots_analysis if screenshot_paths else None,
        additional_context=additional_context
    )

    return result

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
